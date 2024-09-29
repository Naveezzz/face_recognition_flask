from flask import Flask, render_template, request, jsonify, url_for, send_file
import cv2
import numpy as np
from pymongo import MongoClient
import face_recognition
import base64
from datetime import datetime, timedelta
import logging
from bson import ObjectId
import os
from docx import Document  # Import python-docx for Word document creation

# Initialize Flask app
app = Flask(__name__)

# MongoDB Setup
client = MongoClient('mongodb://localhost:27017/')
db = client['attendance_db']
employees_collection = db['employees']
attendance_collection = db['attendance']

# Configure logging
logging.basicConfig(level=logging.INFO)

# Utility function to convert MongoDB ObjectId to string
def object_id_to_str(obj_id):
    return str(obj_id)

# Real-time Face Recognition Route
@app.route('/recognize', methods=['POST'])
def recognize():
    data = request.get_json()
    frame_data = base64.b64decode(data['frame'])
    frame = np.frombuffer(frame_data, dtype=np.uint8)
    frame = cv2.imdecode(frame, cv2.IMREAD_COLOR)

    frame = cv2.resize(frame, (0, 0), fx=0.5, fy=0.5)

    logging.info("Detecting faces in the real-time frame.")
    face_locations = face_recognition.face_locations(frame)
    logging.info("Detected {} faces.".format(len(face_locations)))

    if face_locations:
        face_encodings = face_recognition.face_encodings(frame, face_locations)

        for face_encoding in face_encodings:
            for employee in employees_collection.find():
                if 'face_encoding' not in employee:
                    logging.warning(f"Employee {employee['employee_name']} has no face_encoding. Skipping.")
                    continue

                stored_encoding = np.array(employee['face_encoding'])
                matches = face_recognition.compare_faces([stored_encoding], face_encoding)

                if True in matches:
                    entry_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    attendance_collection.insert_one({
                        'employee_id': employee['employee_id'],
                        'employee_name': employee['employee_name'],
                        'entry_time': entry_time
                    })
                    employee['_id'] = object_id_to_str(employee['_id'])
                    return jsonify({
                        "status": "success",
                        "message": "Successfully Matched",
                        "employee": employee,
                        "entry_time": entry_time
                    })

    return jsonify({"status": "failure", "redirect": url_for('register_page')})

# Registration Page Route
@app.route('/register_page')
def register_page():
    return render_template('register.html')

# Registration Route
@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()
    frame_data = base64.b64decode(data['frame'])
    frame = np.frombuffer(frame_data, dtype=np.uint8)
    frame = cv2.imdecode(frame, cv2.IMREAD_COLOR)

    employee_id = data['employee_id']
    employee_name = data['employee_name']
    entry_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    face_locations = face_recognition.face_locations(frame)
    if len(face_locations) > 0:
        face_encodings = face_recognition.face_encodings(frame, face_locations)
        encoded_face = face_encodings[0]
    else:
        return jsonify({"status": "failure", "message": "No face detected in the frame."})

    employees_collection.insert_one({
        'employee_id': employee_id,
        'employee_name': employee_name,
        'face_encoding': encoded_face.tolist(),
        'last_entry': entry_time
    })

    return jsonify({"status": "success", "message": "Successfully Registered", "redirect": url_for('index')})

# Admin Dashboard Route
@app.route('/admin')
def admin_page():
    return render_template('admin.html')

# Generate Report Route (Updated to trigger Word document download)
@app.route('/generate_report', methods=['POST'])
def generate_report():
    report_type = request.json.get('report_type')
    start_date = None
    end_date = None

    if report_type == 'daily':
        start_date = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = datetime.now().replace(hour=23, minute=59, second=59, microsecond=999999)
    elif report_type == 'weekly':
        start_date = (datetime.now() - timedelta(days=7)).replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = datetime.now().replace(hour=23, minute=59, second=59, microsecond=999999)
    elif report_type == 'monthly':
        start_date = (datetime.now() - timedelta(days=30)).replace(hour=0, minute=0, second=0, microsecond=0)
        end_date = datetime.now().replace(hour=23, minute=59, second=59, microsecond=999999)

    reports = list(attendance_collection.find({
        'entry_time': {
            '$gte': start_date,
            '$lte': end_date
        }
    }))

    if not reports:
        return jsonify({"status": "failure", "message": "No data available for the report."})

    # Create a Word document
    doc = Document()
    doc.add_heading(f'{report_type.capitalize()} Attendance Report', 0)

    # Add table headers
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Employee ID'
    hdr_cells[1].text = 'Employee Name'
    hdr_cells[2].text = 'Entry Time'
    hdr_cells[3].text = 'Attendance ID'

    # Add rows to the table
    for report in reports:
        row_cells = table.add_row().cells
        row_cells[0].text = report.get('employee_id', '')
        row_cells[1].text = report.get('employee_name', '')
        row_cells[2].text = report.get('entry_time', '')
        row_cells[3].text = object_id_to_str(report.get('_id'))

    # Save the Word document in the 'static/reports' folder
    if not os.path.exists('static/reports'):
        os.makedirs('static/reports')

    file_path = f'static/reports/{report_type}_report.docx'
    doc.save(file_path)

    # Use send_file to return the file for download
    return send_file(file_path, as_attachment=True)

# Main Page Route
@app.route('/')
def index():
    return render_template('index.html')

# Run the Flask app
if __name__ == '__main__':
    app.run(debug=True)
